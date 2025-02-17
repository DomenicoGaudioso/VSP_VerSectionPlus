import pandas as pd
import os
import docx
import seaborn as sns
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D  # Per aggiungere custom legend
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
import pyvista as pv
import numpy as np
import ezdxf
from concreteproperties.material import Concrete, SteelBar
from concreteproperties.stress_strain_profile import (
    ConcreteLinear,
    RectangularStressBlock,
    SteelElasticPlastic,
    EurocodeParabolicUltimate,
)
from sectionproperties.pre.library.concrete_sections import concrete_rectangular_section
from concreteproperties.concrete_section import ConcreteSection
from concreteproperties.results import MomentInteractionResults
from concreteproperties.results import BiaxialBendingResults
from sectionproperties.analysis import Section
from sectionproperties.pre import CompoundGeometry, Geometry
from concreteproperties.pre import add_bar
from sklearn.cluster import KMeans
import openpyxl
from openpyxl.drawing.image import Image
from io import BytesIO

def estrai_polinee_dxf(file_dxf, layer_name="SEZ_0"):
    doc = ezdxf.readfile(file_dxf)
    msp = doc.modelspace()
    
    polilinee = []  # Lista per memorizzare le coordinate delle polilinee

    # Analizza tutte le entità presenti nel modelspace
    for entity in msp:
        if entity.dxftype() in ["LWPOLYLINE", "POLYLINE"]:  # Filtra solo polilinee
            if entity.dxf.layer == layer_name:  # Filtra per layer specifico
                punti = [(p[0], p[1]) for p in entity.get_points()]
                polilinee.append(punti)

    return polilinee

def estrai_cerchi_dxf(file_dxf, layer_name=None):
    doc = ezdxf.readfile(file_dxf)
    msp = doc.modelspace()
    
    cerchi = []  # Lista per memorizzare i cerchi trovati

    # Analizza tutte le entità nel modelspace
    for entity in msp:
        if entity.dxftype() == "CIRCLE":  # Se è un cerchio
            if layer_name is None or entity.dxf.layer == layer_name:  # Filtra per layer se specificato
                centro = (entity.dxf.center.x, entity.dxf.center.y)  # Coordinate del centro
                diametro = 2 * entity.dxf.radius  # Calcola il diametro
                cerchi.append((centro, diametro))

    return cerchi

def generate_loft_mesh(sections):
    """Genera i vertici e le facce di una mesh 3D eseguendo un loft tra sezioni."""
    num_sections = len(sections)
    num_points = len(sections[0])
    
    # Creazione delle coordinate dei vertici
    vertices = np.array([point for section in sections for point in section])
    
    # Creazione delle facce
    faces = []
    for i in range(num_sections - 1):
        for j in range(num_points):
            p1 = i * num_points + j
            p2 = i * num_points + (j + 1) % num_points
            p3 = (i + 1) * num_points + j
            p4 = (i + 1) * num_points + (j + 1) % num_points
            
            faces.extend([3, p1, p2, p3, 3, p2, p4, p3])
    
    return vertices, np.array(faces)

import os

def trova_sottocartelle(cartella_principale):
    """
    Trova tutte le sottocartelle dentro la cartella principale e restituisce le loro path.
    """
    if not os.path.exists(cartella_principale):
        print(f"❌ Errore: La cartella '{cartella_principale}' non esiste.")
        return []

    sottocartelle = []

    for nome_cartella in os.listdir(cartella_principale):
        percorso_cartella = os.path.join(cartella_principale, nome_cartella)
        
        if os.path.isdir(percorso_cartella):  # Verifica se è una cartella
            sottocartelle.append(percorso_cartella)

    return sottocartelle

def file_per_estensione(cartella_principale, estensione=".xlsx", iniziali="c"):
    
    """
    Trova tutti i file con una data estensione e che iniziano con una data sequenza di lettere nel nome
    nella cartella principale (senza esplorare sottocartelle).
    """
    if not os.path.exists(cartella_principale):
        print(f"❌ Errore: La cartella '{cartella_principale}' non esiste.")
        return []

    file_trovati = []

    # Esplora solo i file nella cartella principale
    for nome_file in os.listdir(cartella_principale):
        percorso_file = os.path.join(cartella_principale, nome_file)
        
        if os.path.isfile(percorso_file):  # Verifica se è un file
            if nome_file.endswith(estensione) and nome_file.lower().startswith(iniziali.lower()):
                file_trovati.append(percorso_file)

    return file_trovati


#def analizeConcreateSection(pathInfo):

def setmaterial(pathInfo):
    
    for file in os.listdir(pathInfo):
        if file.lower().startswith("cds") and file.lower().endswith(".xlsx"):
            path_excel = os.path.join(pathInfo, file)

    try:
        df = pd.read_excel(path_excel, sheet_name="materiali")

    except FileNotFoundError:
        print(f"❌ Errore: Il file '{path_excel}' non esiste.")

    except ValueError:
        print(f"❌ Errore: Il foglio 'materiali' non è presente nel file.")
        
    
    num_righe = df.shape[0]
    #print("righe", num_righe)
    cls_dict = {}
    steel_dict = {}
    
    for i in range(num_righe):
        
        # DEFINISCO I MATERIALI
        id = int(df.iloc[i]["id"])
        fcd = round(df.iloc[i]["fcd"],2)
        ec2 = df.iloc[i]["ec2"]
        ecu = df.iloc[i]["ecu"]
        fyd = round(df.iloc[0]["fyd"], 2)
        esyd = df.iloc[i]["esyd"]
        esu = df.iloc[i]["esu"]
        
        name_cls = f"Concrete {fcd} MPa" 
        name_steel = f"Steel {fyd} MPa" 
        
        concrete = Concrete(
            name=name_cls,
            density=2.4e-6,
            stress_strain_profile=ConcreteLinear(elastic_modulus=34000),
            ultimate_stress_strain_profile=EurocodeParabolicUltimate(
                compressive_strength=fcd,
                compressive_strain=ec2,
                n=2,
                ultimate_strain=ecu,
            ),
            flexural_tensile_strength=0,
            colour="lightgrey",
        )


        steel = SteelBar(
            name=name_steel,
            density=7.85e-6,
            stress_strain_profile=SteelElasticPlastic(
                yield_strength=fyd,
                elastic_modulus=200e3,
                fracture_strain=esu,
            ),
            colour="red",
        )
        
        cls_dict[id] = concrete
        steel_dict[id] = steel
    
    return cls_dict, steel_dict

def bildSection(pathInfo, cls_dict, steel_dict):

    # COSTRUZIONE DELLA SEZIONE
    
    for file in os.listdir(pathInfo):
        if file.lower().endswith(".dxf"):
            path_dxf = os.path.join(pathInfo, file)

    geometry = estrai_polinee_dxf(path_dxf, layer_name="SEZ_0")
    fcts = [[i, i+1] for i in range(0, len(geometry[0]))]
    fcts[-1] = [len(geometry[0])-1, 0]

    cps = [[0, 0]]
    geom1 = Geometry.from_points(points=geometry[0], facets=fcts, control_points=cps, material=cls_dict[0])
    
    if len(geometry) > 1:
        fcts = [[i, i+1] for i in range(0, len(geometry[1]))]
        fcts[-1] = [len(geometry[1])-1, 0]

        cps = [[0, 0]]
        geom2 = Geometry.from_points(points=geometry[1], facets=fcts, control_points=cps, material=cls_dict[0])
        #geom2.plot_geometry()

    geom = geom1 - geom2
    
    try:
        geometry_add = estrai_polinee_dxf(path_dxf, layer_name="SEZ_1")
        fcts_add = [[i, i+1] for i in range(0, len(geometry_add[0]))]
        fcts_add[-1] = [len(geometry_add[0])-1, 0]
    
        cps_add = [[0, 0]]
        geom3 = Geometry.from_points(points=geometry_add[0], facets=fcts_add, control_points=cps_add, material=cls_dict[1])
        
        geom = geom + (geom3-geom1) 
         
    except:
        print(f"La sezione non ha materiale aggiuntivo")

    #print(len(geometry))
    #geom1.plot_geometry()
    #geom.plot_geometry()

    # Esempio di utilizzo
    layer = "ARM_0"  # Nome del layer (se vuoi tutti i cerchi, metti `None`)
    cerchi_trovati = estrai_cerchi_dxf(path_dxf, layer)

    # Stampa i risultati
    for i, (centro, diametro) in enumerate(cerchi_trovati):
        #print(f"Cerchio {i+1}: Centro {centro}, Diametro {diametro}")
        Area = np.pi*diametro**2/4
        geom = add_bar(geometry=geom, area=Area, material=steel_dict[0], x=centro[0], y=centro[1])

    #geom.plot_geometry(labels=[], cp=False, legend=False)

    conc_sec = ConcreteSection(geom)
    
    return conc_sec

# COSTRUZIONE DELLA SEZIONE
def domino3D(section, cls_dict, steel_dict, cds, n_points = 16, n_level = 10):
    
    ## SOLLECITAZIONI
    Ned = cds["Axial (kN)"]
    Med_y = cds["Moment-y (kN*m)"]
    Med_z = cds["Moment-z (kN*m)"]
    Pcds = np.array([Med_y, Med_z, -Ned]).T
    
    gross_props = section.get_gross_properties()
    #print(steel_dict[0])
    #print(cls_dict[0])
    Nrd_min = gross_props.reinf_lumped_area*steel_dict[0].stress_strain_profile.yield_strength*0.95
    Nrd_max = gross_props.total_area*cls_dict[0].ultimate_stress_strain_profile.compressive_strength + Nrd_min*0.5
    
    
    n_list = np.linspace(-Nrd_min, Nrd_max, n_level)
    biaxial_results = []
    dictDominio = {}

    for n in n_list:
        b_result = conc_sec.biaxial_bending_diagram(n=n, n_points = n_points, progress_bar=False)
        biaxial_results.append(b_result)
        dictDominio[n] = BiaxialBendingResults.get_results_lists(b_result)
        
    PuntiDominio = []
    ## DIAGRAMMA 
    for i in dictDominio:
        x = np.array(dictDominio[i][0])/(1000**2)
        y = np.array(dictDominio[i][1])/(1000**2)
        z = [i/1000]*len(dictDominio[i][0]) 
        
        PuntiDominio.append([(mx, my, nz) for mx, my, nz in zip(x, y, z)])
        
    # Genera la mesh 3D con diametri variabili
    vertices, faces = generate_loft_mesh(PuntiDominio)
    mesh = pv.PolyData(vertices, faces)

    # Assegna una gradazione di colori basata sulla coordinata Z
    scalars = vertices[:, 2]
    Mxrd_max = max(max(vertices[:, 0]), abs(min(vertices[:, 0])))
    Myrd_max = max(max(vertices[:, 1]), abs(min(vertices[:, 1])))
    M_max = max(Mxrd_max, Myrd_max)
    mesh['Elevation'] = scalars
        
    # Definizione del segmento di linea
    # Rendering del risultato
    p = pv.Plotter(off_screen=True)
    p.add_mesh(mesh, show_edges=True, opacity=0.8, cmap="viridis", scalars='Elevation', lighting=False, label="Dominio", show_scalar_bar=False)
    # Esegui il ray tracing
    point_cloud = pv.PolyData(Pcds)
    p.add_mesh(point_cloud, color='yellow', point_size=10.0, render_points_as_spheres=True, label="Sollecitazioni")

    vectors = []
    check = []
    for ip in Pcds:
        start = np.array([0, 0, ip[2]])
        stop = np.array([ip[0]*M_max*1.1, ip[1]*M_max*1.1, ip[2]])
        points, ind = mesh.ray_trace(start, stop, first_point=True)

        # Crea geometrie per visualizzare il ray tracing
        ray = pv.Line(start, stop)
        intersection = pv.PolyData(points) if points.size > 0 else None
        #p.add_mesh(ray, color="blue", line_width=1, label="Ray Segment")
        
        ## PER VEDERE I PUNTI DI INTERSEZIONE
        #if intersection:
            #p.add_mesh(intersection, color="maroon", point_size=10, label="Resistenza", render_points_as_spheres=True)
        
        # Estrarre il primo punto di intersezione
        intersection_point = points if points.size > 0 else [None, None, None]
        vectors.append(intersection_point)
        
        # Calcolo della distanza euclidea
        domanda = np.linalg.norm(start - ip)
        capacita = np.linalg.norm(start - points)
        check.append(domanda/capacita)
        

    vectors =np.array(vectors)

    point_cloud['vectors'] = np.array(vectors)

    arrows = point_cloud.glyph(
        orient='vectors',
        scale=False,
        factor=0.30,
    )

    cds = cds.assign(My_Rd = vectors.T[0]) 
    cds = cds.assign(Mz_Rd = vectors.T[1]) 
    cds = cds.assign(SR = check) 
    cds = cds.drop(['Shear-y (kN)', 'Shear-z (kN)', "Torsion (kN*m)"], axis = 1)


    p.add_mesh(arrows, color='blue')    
    #labels = dict(zlabel='Nrd [KN]', xlabel='Mrd_y [KNm]', ylabel='Mrd_z [KNm]', font_size=10)
    labels = dict(ztitle='Nrd [KN]', xtitle='Mrd_y [KNm]', ytitle='Mrd_z [KNm]', font_size=10)
    p.show_grid(**labels) # Mostrare gli assi
    p.set_background("white") #p.add_legend()
    #p.show(window_size=[2020, 3035]) #, auto_close=False
    #input("Premi Invio per chiudere...")

    
    return p

# 📌 Funzione per salvare il DataFrame come immagine (spezzandolo se necessario)
def save_dataframe_images(df, rows_per_page=30):
    num_parts = (len(df) // rows_per_page) + 1  # Numero di immagini necessarie
    img_paths = []

    for i in range(num_parts):
        start = i * rows_per_page
        end = (i + 1) * rows_per_page
        df_part = df.iloc[start:end]  # Prendi la parte della tabella

        # 🔹 Se la parte del DataFrame è vuota, salta questa iterazione
        if df_part.empty:
            continue

        # 🔹 Arrotonda i valori numerici a due decimali, mantenendo le stringhe invariate
        df_part = df_part.applymap(lambda x: f"{x:.2f}" if isinstance(x, (int, float, np.number)) else x)

        # 🔹 Crea la figura e l'asse
        fig, ax = plt.subplots(figsize=(12, min(1 + len(df_part) * 0.5, 10)))  
        ax.axis("tight")
        ax.axis("off")

        # 🔹 Controlliamo se la colonna "SR" esiste ed è numerica
        if "SR" in df_part.columns:
            sr_values = pd.to_numeric(df_part["SR"], errors="coerce")  # Converte in numeri (NaN se errore)
        else:
            sr_values = pd.Series([0] * len(df_part))  # Se non esiste, valori di default

        # 🔹 Crea la tabella
        table = ax.table(cellText=df_part.values, 
                         colLabels=df_part.columns, 
                         cellLoc="center", 
                         loc="center")

        table.auto_set_font_size(False)
        table.set_fontsize(8)
        table.auto_set_column_width(col=list(range(len(df_part.columns))))  # Adatta larghezza colonne

        # 🔹 Coloriamo le righe con SR > 1 in rosso
        for j, sr in enumerate(sr_values):
            if sr > 1:  # 🔥 SR > 1 → riga in rosso
                for k in range(len(df_part.columns)):
                    table[(j + 1, k)].set_facecolor("red")  # Indice +1 per saltare la riga dell'intestazione

        # 🔹 Salva l'immagine
        img_path = f"table_part_{i}.png"
        plt.savefig(img_path, bbox_inches="tight", dpi=300)
        img_paths.append(img_path)
        plt.close()

    return img_paths

def subplot_figure1(im3d, conc_sec, cls_dict, steel_dict):
    img_pv = im3d.screenshot(return_img=True)
    #pv.close()

    # 📌 Creazione di una figura con gridspec per una disposizione personalizzata
    fig = plt.figure(figsize=(8, 10))
    gs = fig.add_gridspec(3, 2, height_ratios=[2, 1, 1])  # 3 righe, 2 colonne con una riga più alta

    # Aggiungi il primo subplot che occupa tutta la prima riga
    ax1 = fig.add_subplot(gs[0, :])  # Questo subplot occupa tutta la prima riga
    #ax1.set_title("Dominio 3D [Mx - My - N]")

    # Inserisce il plot di PyVista nel primo subplot
    ax1.imshow(img_pv)  # Mostra l'immagine PyVista
    ax1.axis("off")  # Nasconde gli assi

    # Creazione della figura per il secondo grafico (Matplotlib)
    figSec = conc_sec.plot_section()
    fig3 = figSec.figure
    # Converti fig in un'immagine numpy
    canvas = FigureCanvas(fig3)
    canvas.draw()
    img_mpl = np.array(canvas.buffer_rgba())  # Ottieni l'immagine in formato array
    plt.close(fig3)  # Chiudi la figura dopo aver estratto l'immagine

    # Aggiungi il secondo grafico (Matplotlib) nel secondo subplot
    ax2 = fig.add_subplot(gs[1, :])  # Sottogruppo della seconda riga, prima colonna
    ax2.imshow(img_mpl)
    ax2.axis("off")  # Nasconde gli assi

    # Creazione della figura per il terzo grafico (Stress-Strain cls)
    cls_fig = cls_dict[0].ultimate_stress_strain_profile.plot_stress_strain()
    figcls = cls_fig.figure
    # Converti fig in un'immagine numpy
    canvas_cls = FigureCanvas(figcls)
    canvas_cls.draw()
    img_cls = np.array(canvas_cls.buffer_rgba())  # Ottieni l'immagine in formato array
    plt.close(figcls)  # Chiudi la figura dopo aver estratto l'immagine

    # Aggiungi il terzo grafico (Stress-Strain cls) nel terzo subplot
    ax3 = fig.add_subplot(gs[2, 0])  # Sottogruppo della terza riga, prima colonna
    ax3.imshow(img_cls)
    ax3.axis("off")  # Nasconde gli assi

    # Creazione della figura per il quarto grafico (Stress-Strain steel)
    steel_fig = steel_dict[0].stress_strain_profile.plot_stress_strain()
    figsteel = steel_fig.figure
    # Converti fig in un'immagine numpy
    canvas_steel = FigureCanvas(figsteel)
    canvas_steel.draw()
    img_steel = np.array(canvas_steel.buffer_rgba())  # Ottieni l'immagine in formato array
    plt.close(figsteel)  # Chiudi la figura dopo aver estratto l'immagine

    # Aggiungi il quarto grafico (Stress-Strain steel) nel quarto subplot
    ax4 = fig.add_subplot(gs[2, 1])  # Sottogruppo della terza riga, seconda colonna
    ax4.imshow(img_steel)
    ax4.axis("off")  # Nasconde gli assi

    # Modifica i margini
    plt.subplots_adjust(wspace=0.4, hspace=0.2)

    # Mostra il risultato
    plt.tight_layout()
    
    return fig

"""
## CLASTERING DELLE SOLLECITAZIONE
path_cds = r"Z:\studio\CODIFICATE\200 BISantis - due volte prima\Areatecnica\Calcolo\Verifiche Domini\Ritti\cds_Ritto E_sup.xlsx"
#print(path_cds)
cds = pd.read_excel(path_cds, usecols=range(1, 11, 1), skiprows= 1)
Ned = cds["Axial (kN)"]
Med_y = cds["Moment-y (kN*m)"]
Med_z = cds["Moment-z (kN*m)"]


# Reshape dei dati solo per `Ned` (Z)
Ned_reshaped = Ned.values.reshape(-1, 1)  # Convertiamo `Ned` in un array 2D (necessario per K-means)

# Impostiamo il numero di cluster
kmeans = KMeans(n_clusters=10)  # Cambia il numero di cluster se necessario

# Applichiamo K-means solo su `Ned`
kmeans.fit(Ned_reshaped)

# Otteniamo le etichette di clustering per ciascun punto
labels = kmeans.labels_

# Calcoliamo le medie di `Ned` per ogni cluster
cluster_means = []
for i in range(kmeans.n_clusters):
    cluster_points = Ned[labels == i]  # Seleziona i punti che appartengono al cluster i
    cluster_mean = cluster_points.mean()  # Calcola la media
    cluster_means.append(cluster_mean)

# Visualizziamo i risultati
fig = plt.figure()
ax = fig.add_subplot(111, projection='3d')

# Colori diversi per i cluster
scatter = ax.scatter(Med_y, Med_z, Ned, c=labels, cmap='viridis')

# Aggiungiamo le etichette degli assi
ax.set_xlabel('Moment-y (kN*m)')
ax.set_ylabel('Moment-z (kN*m)')
ax.set_zlabel('Axial (kN)')

# Creiamo una legenda personalizzata con le medie dei cluster
legend_labels = [f"Cluster {i+1} (Mean N: {mean:.2f} kN)" for i, mean in enumerate(cluster_means)]
custom_lines = [Line2D([0], [0], marker='o', color='w', markerfacecolor=scatter.cmap(i / len(cluster_means)), markersize=10) for i in range(kmeans.n_clusters)]

# Aggiungiamo la legenda
ax.legend(custom_lines, legend_labels, loc='upper left', fontsize=10)
"""

# Mostra il grafico
#plt.show()

"""
___________________________________________________________________________________________
RUN SCRIPT

"""

pathStar = r"Z:\studio\CODIFICATE\200 BISantis - due volte prima\Areatecnica\Calcolo\Verifiche Domini"
pathSec = trova_sottocartelle(pathStar)
#print(pathSec)

#path = pathSec[1]

word_save = os.path.join(pathStar, "Report Verifiche.docx")
# open an existing document
doc = docx.Document()
doc.add_heading("Report verifiche", level=1)

for i, item in enumerate(pathSec):
    #print(pathSec)
    #print(pathSec[i])
    path = pathSec[i]
    path_cds = file_per_estensione(path, estensione=".xlsx", iniziali="cds")[0]
    #print(path_cds)
    cds = pd.read_excel(path_cds, usecols=range(1, 11, 1), skiprows= 1)
    
    cls_dict, steel_dict = setmaterial(path) # settaggio dei materiali
    conc_sec = bildSection(path, cls_dict, steel_dict) # costruzione della sezione
    im3d = domino3D(conc_sec, cls_dict, steel_dict, cds, n_points=5, n_level=5) # costruzione del dominio 3D
    #im3d.show(interactive=True) #, auto_close=False
    #input("Premi Invio per chiudere...")
    #im3d.screenshot(r"C:\Users\d.gaudioso\Desktop\prova.png", window_size=[2020, 3035])  # Salva l'immagine in un file PNG
    figure = subplot_figure1(im3d, conc_sec, cls_dict, steel_dict)
    image_stream = BytesIO()
    figure.savefig(image_stream, format="png")  # Salva l'immagine in memoria
    image_stream.seek(0)  # Torna all'inizio del file in memoria

    figure.savefig(path)
    print('Figure written File successfully.')

    # saving the excel
    nomefile = "verifiche_MxMyN_" + os.path.basename(path) +".xlsx"
    cds.to_excel(os.path.join(path, nomefile))
    print('DataFrame is written to Excel File successfully.')

    ### TO WORD##
    doc.add_heading(os.path.basename(path), level=2)
    #Inserimento dell'immagine dal buffer di memoria
    doc.add_picture(image_stream, width=docx.shared.Inches(5))  # Inserisce l'immagine dal buffer

    # 📌 Generiamo le immagini del dataframe
    image_paths = save_dataframe_images(cds, rows_per_page=70)

    # 📌 Inseriamo le immagini in Word
    doc.add_paragraph("Risultati in forma tabellare")

    for img in image_paths:
        doc.add_picture(img, width = docx.shared.Inches(6))
        doc.add_page_break()  # Aggiunge un'interruzione di pagina

# 📌 Salviamo il documento
doc.save(word_save)